# https://python-docx.readthedocs.org
import docx
import os

os.chdir('docs')

d = docx.Document('demo.docx')
print(d.paragraphs) # objects
print(d.paragraphs[0])
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)
p = d.paragraphs[1]
print(p.runs) # [obj] -> changes
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[1].bold) # bool
print(p.runs[2].text)
print(p.runs[3].italic) # bool
print(p.runs[3].text)

p.runs[3].underline = True
p.runs[3].text = 'italic and underline'
#d.save('demo2.docx')

print(p.style)
p.style = 'Title'
#d.save('demo3.docx')

# New document
d1 = docx.Document()
d1.add_paragraph('Hello this is a paragraph.')
d1.add_paragraph('Hello this is another paragraph.')
#d1.save('demo4.docx')
p1 = d1.paragraphs[0]
p1.add_run('This is a new run')
p1.runs[1].bold = True
d1.save('demo5.docx')









